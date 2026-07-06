"""DOCX report. Requires python-docx; falls back to Markdown otherwise."""
from __future__ import annotations

import datetime as _dt
from collections.abc import Iterable
from io import BytesIO

from ._common import (
    SEV_ORDER,
    coverage_rows,
    coverage_rows_by_host,
    curl_from_reproduction,
    reqlore_version,
    utc_now,
)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    Document = None  # type: ignore[assignment]
    DOCX_AVAILABLE = False


_SEV_COLOR = {
    "critical": RGBColor(0xA5, 0x1A, 0x1A) if DOCX_AVAILABLE else None,
    "high":     RGBColor(0xC9, 0x3B, 0x00) if DOCX_AVAILABLE else None,
    "medium":   RGBColor(0x9A, 0x6A, 0x00) if DOCX_AVAILABLE else None,
    "low":      RGBColor(0x2A, 0x6E, 0x2A) if DOCX_AVAILABLE else None,
    "info":     RGBColor(0x23, 0x4E, 0x7A) if DOCX_AVAILABLE else None,
}


def render_docx(project_meta: dict, findings: Iterable[dict], *,
                 title: str = "Reqlore Security Findings",
                 now: _dt.datetime | None = None,
                 classification: str = "",
                 include_coverage: bool = False,
                 coverage: Iterable[dict] | None = None,
                 coverage_by_host: Iterable[dict] | None = None,
                 reproductions: dict[str, dict] | None = None) -> bytes:
    """Return the .docx file as bytes. Raises RuntimeError if python-docx is
    missing — callers should check ``DOCX_AVAILABLE`` first or fall back to
    :func:`render_markdown`.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed. Install it with "
            "'pip install python-docx' or export the report as Markdown / HTML "
            "instead."
        )
    findings = list(findings)
    ts = utc_now(now).isoformat(timespec="seconds")
    version = reqlore_version()
    doc = Document()
    doc.add_heading(title, level=0)
    if classification:
        banner = doc.add_paragraph()
        run = banner.add_run(classification.upper())
        run.bold = True
    p = doc.add_paragraph()
    p.add_run(f"Project: {project_meta.get('name', '?')}").bold = True
    p.add_run("    Generated (UTC): ")
    p.add_run(ts).bold = True
    p.add_run("    Total findings: ")
    p.add_run(str(len(findings))).bold = True

    # Summary table
    doc.add_heading("Summary", level=1)
    counts = dict.fromkeys(SEV_ORDER, 0)
    for f in findings:
        counts[f["severity"]] = counts.get(f["severity"], 0) + 1
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Severity"
    hdr[1].text = "Count"
    for sev in SEV_ORDER:
        row = table.add_row().cells
        row[0].text = sev.title()
        row[1].text = str(counts[sev])

    if include_coverage:
        rows = coverage_rows(coverage)
        doc.add_heading("Coverage", level=1)
        if not rows:
            doc.add_paragraph("No rule runs recorded.")
        else:
            ctable = doc.add_table(rows=1, cols=3)
            ctable.style = "Light Grid"
            chdr = ctable.rows[0].cells
            chdr[0].text = "Rule"
            chdr[1].text = "Fired"
            chdr[2].text = "Evaluated"
            for r in rows:
                row = ctable.add_row().cells
                row[0].text = r["rule_id"]
                row[1].text = str(r["fired"])
                row[2].text = str(r["evaluated"])
        per_host = coverage_rows_by_host(coverage_by_host)
        if per_host:
            doc.add_heading("Coverage by host", level=2)
            htable = doc.add_table(rows=1, cols=4)
            htable.style = "Light Grid"
            hhdr = htable.rows[0].cells
            hhdr[0].text = "Rule"
            hhdr[1].text = "Host"
            hhdr[2].text = "Fired"
            hhdr[3].text = "Evaluated"
            for r in per_host:
                row = htable.add_row().cells
                row[0].text = r["rule_id"]
                row[1].text = r["host"]
                row[2].text = str(r["fired"])
                row[3].text = str(r["evaluated"])

    # Per-severity sections
    for sev in SEV_ORDER:
        bucket = [f for f in findings if f["severity"] == sev]
        if not bucket:
            continue
        doc.add_heading(f"{sev.title()} ({len(bucket)})", level=1)
        for f in bucket:
            heading = doc.add_heading(f["title"], level=2)
            for run in heading.runs:
                run.font.color.rgb = _SEV_COLOR[sev]
            meta_bits: list[str] = []
            if f.get("rule_id"):
                meta_bits.append(f"Rule: {f['rule_id']}")
            if f.get("source"):
                meta_bits.append(f"Source: {f['source']}")
            if f.get("host"):
                meta_bits.append(f"Host: {f['host']}")
            if f.get("url"):
                meta_bits.append(f"URL: {f['url']}")
            if f.get("cwe"):
                meta_bits.append(f"CWE: {f['cwe']}")
            if f.get("owasp"):
                meta_bits.append(f"OWASP: {f['owasp']}")
            if f.get("cvss_score") not in (None, ""):
                vec = f.get("cvss_vector") or ""
                meta_bits.append(
                    f"CVSS: {f['cvss_score']}" + (f" ({vec})" if vec else "")
                )
            _conf = f.get("confidence") or "firm"
            meta_bits.append(f"Confidence: {_conf}")
            _occ = f.get("occurrence_count") or 1
            if _occ and _occ > 1:
                meta_bits.append(f"Occurrences: {int(_occ)}")
            _tags = f.get("fingerprint_tags_list") or (
                [t for t in (f.get("fingerprint_tags") or "").split(",") if t]
            )
            if _tags:
                meta_bits.append("Response signals: " + ", ".join(_tags))
            meta_bits.append(f"Status: {f.get('status', 'open')}")
            doc.add_paragraph("    ·  ".join(meta_bits))
            if f.get("description"):
                doc.add_paragraph("Description:").runs[0].bold = True
                doc.add_paragraph(str(f["description"]))
            if f.get("evidence"):
                doc.add_paragraph("Evidence:").runs[0].bold = True
                ev = doc.add_paragraph(_clip(f["evidence"], 1500))
                ev.runs[0].font.name = "Consolas"
                ev.runs[0].font.size = Pt(9)
            if f.get("payload"):
                doc.add_paragraph("Payload:").runs[0].bold = True
                pl = doc.add_paragraph(_clip(f["payload"], 500))
                pl.runs[0].font.name = "Consolas"
                pl.runs[0].font.size = Pt(9)
            curl = _curl_for(f, reproductions)
            if curl:
                doc.add_paragraph("Reproduction:").runs[0].bold = True
                rep = doc.add_paragraph(curl)
                rep.runs[0].font.name = "Consolas"
                rep.runs[0].font.size = Pt(9)
            if f.get("remediation"):
                doc.add_paragraph("Remediation:").runs[0].bold = True
                doc.add_paragraph(str(f["remediation"]))
            refs = f.get("references") or []
            if refs:
                doc.add_paragraph("References:").runs[0].bold = True
                for ref in refs:
                    doc.add_paragraph(str(ref), style="List Bullet")

    footer = doc.add_paragraph()
    footer.add_run(f"Generated by reqlore {version} at {ts}.").italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _curl_for(finding: dict, reproductions: dict[str, dict] | None) -> str:
    if not reproductions:
        return ""
    token = finding.get("reproduction_token")
    if not token:
        return ""
    return curl_from_reproduction(reproductions.get(token))


def _clip(s: str, n: int) -> str:
    if len(s) <= n:
        return s
    return s[:n] + f"\n... ({len(s) - n} more chars)"
